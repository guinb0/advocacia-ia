"""Ingere o conhecimento trabalhista versionado e os modelos DOCX no pgvector.

Idempotente por caminho: arquivo inalterado é ignorado; arquivo alterado substitui
atomicamente sua fonte e seus chunks somente depois de os embeddings ficarem prontos.
"""

from __future__ import annotations

import hashlib
import json
import os
import time
from dataclasses import dataclass
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree

import psycopg
from docx import Document

from app.rag import carregar_env, gerar_embeddings, vetor_literal
from scripts.ingerir_jurimetria import dividir, limpar_texto

BASE = Path(__file__).resolve().parent.parent
WORKSPACE = BASE.parent.parent
LEGAL = WORKSPACE / "ia-juridica" / "src" / "legal_agent" / "legal"
DOCS = BASE / "docs"
DOCX_TRABALHISTAS = (
    "CHECK LIST ACIDENTE DO TRABALHO 31.07.26.docx",
    "CHECK LIST ASSALTO.docx",
    "CHECK LIST DOENÇA OCUPACIONAL.docx",
    "CONTRATO oficial.docx",
    "DECLARACAO DE HIPOSSUFICIENCIA ECONOMICA.docx",
    "PROCURACAO.docx",
)
CONEXAO = {
    "connect_timeout": 15,
    "keepalives": 1,
    "keepalives_idle": 30,
    "keepalives_interval": 10,
    "keepalives_count": 3,
    "options": "-c statement_timeout=180000 -c lock_timeout=10000",
    "tcp_user_timeout": 30000,
}


@dataclass(frozen=True)
class ArquivoJuridico:
    caminho: Path
    relativo: str
    categoria: str

    @property
    def identificador(self) -> str:
        return f"direito-trabalhista:{self.relativo.lower()}"


def arquivos() -> list[ArquivoJuridico]:
    encontrados = [
        ArquivoJuridico(p, f"legal/{p.relative_to(LEGAL).as_posix()}", p.parent.name)
        for p in LEGAL.rglob("*")
        if p.is_file() and p.suffix.lower() in {".yaml", ".yml"}
    ]
    encontrados.extend(
        ArquivoJuridico(DOCS / nome, f"docs/{nome}", "documento_escritorio")
        for nome in DOCX_TRABALHISTAS
        if (DOCS / nome).is_file()
    )
    return sorted(encontrados, key=lambda item: item.relativo.casefold())


def texto_docx(caminho: Path) -> str:
    doc = Document(caminho)
    partes = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for tabela in doc.tables:
        for linha in tabela.rows:
            celulas = [" ".join(c.text.split()) for c in linha.cells]
            if any(celulas):
                partes.append(" | ".join(celulas))
    texto = limpar_texto("\n".join(partes))
    # Alguns modelos antigos guardam tudo em caixas de texto, invisíveis para
    # python-docx. O XML ainda contém cada `w:t`; sem este fallback um DOCX
    # aparentemente normal entraria como documento vazio.
    if len(texto) < 80:
        with ZipFile(caminho) as pacote:
            raiz = ElementTree.fromstring(pacote.read("word/document.xml"))
        partes_xml = [no.text.strip() for no in raiz.iter() if no.tag.endswith("}t") and no.text and no.text.strip()]
        texto = limpar_texto("\n".join(partes_xml))
    return texto


def extrair_texto(caminho: Path) -> str:
    if caminho.suffix.lower() == ".docx":
        return texto_docx(caminho)
    return limpar_texto(caminho.read_text(encoding="utf-8"))


def _embeddings_com_retry(chunks: list[str], tentativas: int = 5) -> list[list[float]]:
    for tentativa in range(1, tentativas + 1):
        try:
            return gerar_embeddings(chunks, timeout=180)
        except Exception:
            if tentativa == tentativas:
                raise
            time.sleep(min(2 ** tentativa, 20))
    raise AssertionError("inalcançável")


def _conectar_com_retry(tentativas: int = 12) -> psycopg.Connection:
    ultimo: Exception | None = None
    for tentativa in range(1, tentativas + 1):
        try:
            return psycopg.connect(os.environ["DATABASE_URL"], autocommit=True, **CONEXAO)
        except (psycopg.OperationalError, psycopg.errors.ConnectionTimeout) as exc:
            ultimo = exc
            if tentativa < tentativas:
                time.sleep(min(tentativa * 2, 15))
    assert ultimo is not None
    raise ultimo


def _sha_existente(identificador: str) -> str | None:
    with _conectar_com_retry() as banco:
        linha = banco.execute(
            """SELECT k.metadados->>'sha256'
                 FROM fontes f LEFT JOIN knowledge_chunks k ON k.fonte_id=f.id
                WHERE f.tipo='interno' AND f.identificador=%s
                ORDER BY k.ordem LIMIT 1""",
            (identificador,),
        ).fetchone()
        return linha[0] if linha else None


def _gravar(
    item: ArquivoJuridico, sha: str, chunks: list[str], vetores: list[list[float]]
) -> str:
    """Grava em conexão curta; retry posterior pode repetir sem duplicar."""
    metadados = json.dumps(
        {
            "origem": "acervo_interno",
            "dominio": "direito_trabalho",
            "categoria": item.categoria,
            "arquivo": item.relativo,
            "sha256": sha,
            "tipo_documento": item.caminho.suffix.lower().lstrip("."),
        },
        ensure_ascii=False,
    )
    ultimo: Exception | None = None
    for tentativa in range(1, 13):
        try:
            with _conectar_com_retry() as banco, banco.transaction():
                existente = banco.execute(
                    "SELECT id FROM fontes WHERE tipo='interno' AND identificador=%s FOR UPDATE",
                    (item.identificador,),
                ).fetchone()
                if existente:
                    fonte_id = existente[0]
                    atual = banco.execute(
                        "SELECT metadados->>'sha256' FROM knowledge_chunks WHERE fonte_id=%s ORDER BY ordem LIMIT 1",
                        (fonte_id,),
                    ).fetchone()
                    if atual and atual[0] == sha:
                        return "ignorado"
                    banco.execute("DELETE FROM knowledge_chunks WHERE fonte_id=%s", (fonte_id,))
                    banco.execute("UPDATE fontes SET titulo=%s WHERE id=%s", (item.caminho.stem, fonte_id))
                    acao = "atualizado"
                else:
                    fonte_id = banco.execute(
                        """INSERT INTO fontes(tipo,titulo,identificador)
                           VALUES ('interno',%s,%s) RETURNING id""",
                        (item.caminho.stem, item.identificador),
                    ).fetchone()[0]
                    acao = "inserido"
                with banco.cursor() as cursor:
                    cursor.executemany(
                        """INSERT INTO knowledge_chunks(fonte_id,ordem,texto,metadados,embedding)
                           VALUES (%s,%s,%s,%s::jsonb,%s::vector)""",
                        [
                            (fonte_id, ordem, chunk, metadados, vetor_literal(vetor))
                            for ordem, (chunk, vetor) in enumerate(zip(chunks, vetores, strict=True))
                        ],
                    )
                return acao
        except (psycopg.OperationalError, psycopg.errors.QueryCanceled) as exc:
            ultimo = exc
            if tentativa < 12:
                time.sleep(min(tentativa * 2, 15))
    assert ultimo is not None
    raise ultimo


def ingerir() -> dict[str, int]:
    carregar_env()
    itens = arquivos()
    stats = {"arquivos": len(itens), "inseridos": 0, "atualizados": 0, "ignorados": 0, "chunks": 0}
    for indice, item in enumerate(itens, 1):
        sha = hashlib.sha256(item.caminho.read_bytes()).hexdigest()
        if _sha_existente(item.identificador) == sha:
            stats["ignorados"] += 1
            print(f"[{indice}/{len(itens)}] já existe: {item.relativo}", flush=True)
            continue
        texto = extrair_texto(item.caminho)
        chunks = dividir(texto)
        if not chunks:
            print(f"[{indice}/{len(itens)}] sem texto útil: {item.relativo}", flush=True)
            continue
        # A conexão fica fechada durante esta chamada externa potencialmente lenta.
        vetores = _embeddings_com_retry(chunks)
        acao = _gravar(item, sha, chunks, vetores)
        if acao == "ignorado":
            stats["ignorados"] += 1
        elif acao == "atualizado":
            stats["atualizados"] += 1
        else:
            stats["inseridos"] += 1
        stats["chunks"] += len(chunks)
        print(f"[{indice}/{len(itens)}] {item.relativo}: {len(chunks)} chunks ({acao})", flush=True)
    return stats


if __name__ == "__main__":
    for tentativa in range(1, 31):
        try:
            print(ingerir())
            break
        except Exception as exc:
            print(f"tentativa {tentativa}/30 falhou: {type(exc).__name__}: {str(exc)[:160]}", flush=True)
            if tentativa == 30:
                raise
            time.sleep(10)
